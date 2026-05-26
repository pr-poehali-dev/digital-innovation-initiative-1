"""
Экспорт результата генерации в DOCX (Word).
Для дипломов, рефератов, эссе, аналитических записок.
"""
import json
import os
import re
import base64
import io
import psycopg2


def get_db():
    conn = psycopg2.connect(os.environ["DATABASE_URL"])
    conn.autocommit = False
    return conn


def get_schema():
    return os.environ.get("MAIN_DB_SCHEMA", "public")


ALLOWED_ORIGINS = {
    "https://raven.moscow",
    "https://www.raven.moscow",
    "https://docmind.ai",
    "https://digital-innovation-initiative-1--preview.poehali.dev",
    "https://poehali.dev",
    "http://localhost:5173",
    "http://localhost:3000",
}


def _is_allowed_origin(origin: str) -> bool:
    """Безопасная проверка origin через urlparse — не raw endswith.
    Защита от попыток типа https://attacker.com/.poehali.dev"""
    if not origin:
        return False
    if origin in ALLOWED_ORIGINS:
        return True
    try:
        from urllib.parse import urlparse
        parsed = urlparse(origin)
        # Только https + допустимый hostname
        if parsed.scheme not in ("https", "http"):
            return False
        hostname = (parsed.hostname or "").lower()
        # Точное совпадение с poehali.dev или его поддоменом
        if hostname == "poehali.dev" or hostname.endswith(".poehali.dev"):
            return True
        return False
    except Exception:
        return False


def cors_headers(origin: str = None):
    """Strict CORS: deny-by-default. Без Allow-Credentials — auth через header, не cookies."""
    headers = {
        "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type, X-Session-Id",
        "Vary": "Origin",
    }
    if _is_allowed_origin(origin):
        headers["Access-Control-Allow-Origin"] = origin
    return headers


def json_response(data, status=200, origin=None):
    return {
        "statusCode": status,
        "headers": {**cors_headers(origin), "Content-Type": "application/json"},
        "body": json.dumps(data, ensure_ascii=False, default=str),
    }


def get_current_user(conn, session_id):
    if not session_id:
        return None
    schema = get_schema()
    cur = conn.cursor()
    cur.execute(
        f"SELECT u.id, u.name FROM {schema}.sessions s JOIN {schema}.users u ON u.id = s.user_id WHERE s.id = %s AND s.expires_at > NOW()",
        (session_id,),
    )
    row = cur.fetchone()
    return {"id": row[0], "name": row[1]} if row else None


def build_docx(content: str, title: str, author: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Поля по ГОСТ
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Стиль по умолчанию
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(1.25)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Титульная страница
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True

    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"\nАвтор: {author}")
        author_run.font.size = Pt(12)
        author_run.font.italic = True

    doc.add_page_break()

    # Парсим контент построчно
    lines = content.split("\n")
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue

        # Заголовки уровня 1: # Заголовок или ВСЕ ЗАГЛАВНЫМИ
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.paragraph_format.first_line_indent = Cm(0)
            continue
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.paragraph_format.first_line_indent = Cm(0)
            continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            h.paragraph_format.first_line_indent = Cm(0)
            continue

        # Слайд N: → как подзаголовок
        if re.match(r"^Слайд\s+\d+[:.]", line, re.IGNORECASE):
            h = doc.add_heading(line, level=2)
            h.paragraph_format.first_line_indent = Cm(0)
            continue

        # Жирный текст **Заголовок**
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(line.strip("*"))
            run.bold = True
            continue

        # Буллеты
        if re.match(r"^[•\-\*\–\—]\s", line) or re.match(r"^\d+\.\s", line):
            clean = re.sub(r"^[•\-\*\–\—]\s|^\d+\.\s", "", line).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            # Убираем markdown bold внутри
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", clean)
            p.add_run(clean)
            continue

        # Обычный абзац
        clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        doc.add_paragraph(clean)

    # Сохраняем в байты
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def handler(event: dict, context) -> dict:
    origin = (event.get("headers") or {}).get("Origin") or (event.get("headers") or {}).get("origin")

    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": cors_headers(origin), "body": ""}

    method = event.get("httpMethod", "GET")
    body = {}
    if event.get("body"):
        try:
            body = json.loads(event["body"])
        except Exception:
            pass

    session_id = event.get("headers", {}).get("X-Session-Id", "")
    conn = get_db()
    schema = get_schema()

    try:
        user = get_current_user(conn, session_id)
        if not user:
            return json_response({"error": "Не авторизован"}, 401, origin=origin)

        if method != "POST":
            return json_response({"error": "Method not allowed"}, 405, origin=origin)

        action = body.get("action", "export_run")

        # ── audit_report: DOCX-отчёт прозрачного аудита ──────────────────
        if action == "audit_report":
            audit_id = body.get("audit_id")
            if not audit_id:
                return json_response({"error": "Нужен audit_id"}, 400, origin=origin)

            cur = conn.cursor()
            cur.execute(
                f"""SELECT result_json, source_filename FROM {schema}.audit_runs
                    WHERE id = %s AND user_id = %s""",
                (int(audit_id), user["id"]),
            )
            row = cur.fetchone()
            if not row:
                return json_response({"error": "Аудит не найден"}, 404, origin=origin)

            result_raw, source_filename = row
            audit = json.loads(result_raw) if result_raw else {}
            summary = audit.get("audit_summary") or {}
            findings = audit.get("findings") or []
            criteria = audit.get("criteria") or []
            compliance = audit.get("compliance_matrix") or []
            unverified = audit.get("unverified_items") or []
            documents_used = audit.get("documents_used") or []
            presentation_name = source_filename or "Презентация"

            # Считаем цифры из реальных массивов, не из summary AI
            real_findings = len(findings)
            real_matched = len([c for c in compliance if c.get("status") == "met"])
            real_partial = len([c for c in compliance if c.get("status") == "partially_met"])
            real_not_met = len([c for c in compliance if c.get("status") == "not_met"])
            real_unverified = len(unverified)

            from datetime import datetime
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io as _io

            def add_heading(doc, text, level=1):
                h = doc.add_heading(text, level=level)
                h.paragraph_format.first_line_indent = Cm(0)
                return h

            def add_para(doc, text, bold=False, italic=False, color=None):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                if color:
                    run.font.color.rgb = RGBColor(*color)
                return p

            def add_kv(doc, key, value):
                if not value:
                    return
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                r1 = p.add_run(f"{key}: ")
                r1.bold = True
                p.add_run(str(value))

            doc = Document()
            for section in doc.sections:
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(1.5)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)

            # Титул
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr = tp.add_run("Отчёт аудита презентации")
            tr.font.size = Pt(18)
            tr.font.bold = True

            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.add_run(presentation_name).font.size = Pt(13)

            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dp.add_run(datetime.now().strftime("%d.%m.%Y")).font.size = Pt(11)
            doc.add_page_break()

            # 1. Сводка
            add_heading(doc, "1. Сводка результатов", 1)
            score = summary.get("compliance_score")
            add_kv(doc, "Файл", presentation_name)
            add_kv(doc, "Слайдов проверено", audit.get("slide_count") or summary.get("total_slides"))
            add_kv(doc, "Документов использовано", len(documents_used) or audit.get("document_count"))
            add_kv(doc, "Замечаний найдено", real_findings)
            add_kv(doc, "Критериев проверено", len(compliance))
            if score is not None:
                add_kv(doc, "Оценка соответствия", f"{score}%")
            add_kv(doc, "Соответствует критериям", real_matched)
            add_kv(doc, "Частичное соответствие", real_partial)
            add_kv(doc, "Не соответствует", real_not_met)
            add_kv(doc, "Не удалось проверить", real_unverified)

            key_risks = summary.get("key_risks") or []
            if key_risks:
                add_para(doc, "Ключевые риски:", bold=True)
                for r in key_risks:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.add_run(r)

            # 2. Использованные документы
            if documents_used:
                doc.add_page_break()
                add_heading(doc, "2. Использованные документы", 1)
                role_label_map = {
                    "standard": "Стандарт", "criteria": "Критерии",
                    "source": "Источник", "template": "Шаблон", "material": "Материал"
                }
                for d in documents_used:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.first_line_indent = Cm(0)
                    role_txt = role_label_map.get(d.get("role",""), d.get("role",""))
                    p.add_run(f"{d.get('name','')} — [{role_txt}]")

            # 3. Извлечённые критерии
            if criteria:
                doc.add_page_break()
                add_heading(doc, "3. Критерии проверки", 1)
                add_para(doc, f"Из документов извлечено {len(criteria)} критериев:", italic=True)
                for cr in criteria:
                    doc.add_paragraph()
                    add_para(doc, f"[{cr.get('criterion_id','?')}] {cr.get('title','')}", bold=True)
                    add_kv(doc, "Роль", cr.get("role"))
                    add_kv(doc, "Источник", cr.get("source_document"))
                    if cr.get("description"):
                        add_para(doc, cr["description"])
                    if cr.get("source_quote"):
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Cm(0)
                        p.add_run(f"«{cr['source_quote']}»").italic = True

            # 3. Замечания
            if findings:
                doc.add_page_break()
                add_heading(doc, "3. Замечания и несоответствия", 1)
                sev_label = {"critical": "Критично", "high": "Высокий", "medium": "Средний", "low": "Низкий"}
                for i, f in enumerate(findings, 1):
                    doc.add_paragraph()
                    add_para(doc, f"Замечание {i}: {f.get('short_title','')}", bold=True)
                    sev = sev_label.get(f.get("severity",""), f.get("severity",""))
                    add_kv(doc, "Приоритет", sev)
                    add_kv(doc, "Слайд", f"{f.get('slide_index','?')} — {f.get('slide_title','')}")
                    add_kv(doc, "Критерий", f.get("violated_criterion") or f.get("issue_type",""))
                    add_kv(doc, "Уверенность системы", f.get("confidence",""))
                    if f.get("explanation"):
                        add_para(doc, f.get("explanation",""))
                    if f.get("what_required"):
                        add_para(doc, "Что требовал документ:", bold=True)
                        add_para(doc, f.get("what_required",""))
                    if f.get("what_found"):
                        add_para(doc, "Что найдено в презентации:", bold=True)
                        add_para(doc, f.get("what_found",""))
                    if f.get("gap_description"):
                        add_para(doc, "В чём расхождение:", bold=True)
                        add_para(doc, f.get("gap_description",""))
                    if f.get("evidence_from_presentation"):
                        add_kv(doc, "Цитата из презентации", f"«{f['evidence_from_presentation']}»")
                    if f.get("evidence_from_source_docs"):
                        add_kv(doc, f"Цитата из {f.get('related_document_name','документа')}", f"«{f['evidence_from_source_docs']}»")
                    if f.get("suggested_fix"):
                        add_para(doc, "Рекомендация:", bold=True)
                        add_para(doc, f.get("suggested_fix",""))

            # 4. Матрица соответствия
            if compliance:
                doc.add_page_break()
                add_heading(doc, "4. Матрица соответствия критериям", 1)
                status_label = {"met": "Соответствует", "partially_met": "Частично", "not_met": "Не соответствует", "not_checked": "Не проверено"}
                for c in compliance:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.first_line_indent = Cm(0)
                    sl = status_label.get(c.get("status",""), c.get("status",""))
                    slide_txt = f" (Слайд {c['slide_index']})" if c.get("slide_index") else ""
                    p.add_run(f"[{sl}] {c.get('criterion','')} — {c.get('source','')}{slide_txt}")
                    if c.get("comment"):
                        p2 = doc.add_paragraph()
                        p2.paragraph_format.left_indent = Cm(1)
                        p2.paragraph_format.first_line_indent = Cm(0)
                        p2.add_run(c["comment"]).italic = True

            # 5. Что не удалось проверить
            if unverified:
                doc.add_page_break()
                add_heading(doc, "5. Что система не смогла проверить", 1)
                add_para(doc, "Следующие критерии требуют ручной проверки:", italic=True)
                reason_label = {
                    "insufficient_data": "Недостаточно данных",
                    "ambiguous_criterion": "Неоднозначный критерий",
                    "missing_section": "Раздел отсутствует",
                    "no_relevant_slide": "Нет подходящего слайда",
                }
                for u in unverified:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.first_line_indent = Cm(0)
                    reason = reason_label.get(u.get("reason",""), u.get("reason",""))
                    p.add_run(f"{u.get('criterion','')} — {reason}")
                    if u.get("reason_text"):
                        p2 = doc.add_paragraph()
                        p2.paragraph_format.left_indent = Cm(1)
                        p2.paragraph_format.first_line_indent = Cm(0)
                        p2.add_run(u["reason_text"]).italic = True

            buf = _io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()
            b64 = base64.b64encode(docx_bytes).decode("utf-8")
            safe_name = re.sub(r"[^\w\d\-_\u0400-\u04FF]", "_", (presentation_name or "audit")[:40])

            return json_response({
                "filename": f"Аудит_{safe_name}.docx",
                "file_data": b64,
            }, origin=origin)

        # ── export_run: оригинальная логика ──────────────────────────────
        run_id = body.get("run_id")
        if not run_id:
            return json_response({"error": "Нужен run_id"}, 400, origin=origin)

        cur = conn.cursor()
        cur.execute(
            f"""SELECT gr.result_json, gr.task_id FROM {schema}.generation_runs gr WHERE gr.id = %s""",
            (run_id,),
        )
        row = cur.fetchone()
        if not row:
            return json_response({"error": "Результат не найден"}, 404, origin=origin)

        result_json, task_id = row
        cur.execute(
            f"SELECT title, topic, project_id FROM {schema}.tasks WHERE id = %s",
            (task_id,),
        )
        task_row = cur.fetchone()
        if not task_row:
            return json_response({"error": "Задание не найдено"}, 404, origin=origin)

        task_title, task_topic, project_id = task_row
        cur.execute(
            f"SELECT role FROM {schema}.project_members WHERE project_id = %s AND user_id = %s",
            (project_id, user["id"]),
        )
        if not cur.fetchone():
            return json_response({"error": "Нет доступа"}, 403, origin=origin)

        content = ""
        if result_json:
            try:
                content = json.loads(result_json).get("content", "")
            except Exception:
                content = result_json

        if not content:
            return json_response({"error": "Нет контента для экспорта"}, 400, origin=origin)

        doc_title = task_topic or task_title or "Работа"
        docx_bytes = build_docx(content, doc_title, user.get("name", ""))
        b64 = base64.b64encode(docx_bytes).decode("utf-8")
        safe_name = re.sub(r"[^\w\d\-_\u0400-\u04FF]", "_", doc_title[:50])

        return json_response({
            "filename": f"{safe_name}.docx",
            "file_data": b64,
        }, origin=origin)

    finally:
        conn.close()